"""Word export for printable TextCore course material."""

from __future__ import annotations

import html
import io
import re
from collections.abc import Iterable
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from textcore.contracts.course_state import DEFAULT_VERSION, VERSION_KEYS

DEFAULT_EXPORT_SECTIONS = ("summary", "concise", "cards", "materials", "review")
EXPORT_FORMATS = {"printable", "archive"}
VERSION_LABELS = {
    "faithful": "保真清洗",
    "concise": "精简整理",
    "study": "学习整理",
    "outline": "结构提纲",
}
REVIEW_GRAY = RGBColor(104, 104, 104)


def export_course_docx(
    course_state: dict[str, Any],
    *,
    sections: Iterable[str] | None = None,
    export_format: str = "printable",
    current_version: str | None = None,
) -> bytes:
    """Render a course_state to a real .docx file."""

    selected = normalize_sections(sections, current_version=current_version)
    doc = Document()
    _configure_document(doc, export_format)

    _add_front_matter(doc, course_state, export_format)

    if "summary" in selected:
        _add_summary(doc, course_state)
    for version_key in VERSION_KEYS:
        if version_key in selected:
            _add_version(doc, version_key, course_state.get("versions", {}).get(version_key, {}))
    if "cards" in selected:
        _add_cards(doc, course_state.get("knowledge_cards", []))
    if "materials" in selected:
        _add_materials(doc, course_state.get("writing_materials", []))
    if "review" in selected:
        _add_review_flags(doc, course_state.get("review_flags", []), export_format)
    if "classics" in selected:
        _add_classics(doc, course_state.get("classics_refs", []), export_format)
    if export_format == "archive":
        _add_archive_notes(doc, course_state)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def normalize_sections(
    sections: Iterable[str] | None,
    *,
    current_version: str | None = None,
) -> set[str]:
    raw_sections = list(sections or DEFAULT_EXPORT_SECTIONS)
    normalized: set[str] = set()
    selected_current = current_version if current_version in VERSION_KEYS else DEFAULT_VERSION
    aliases = {
        "course_summary": "summary",
        "knowledge_cards": "cards",
        "writing_materials": "materials",
        "review_flags": "review",
        "classics_refs": "classics",
    }
    for section in raw_sections:
        key = aliases.get(section, section)
        if key == "current_version":
            normalized.add(selected_current)
        elif key in {*VERSION_KEYS, "summary", "cards", "materials", "review", "classics"}:
            normalized.add(key)
    return normalized or set(DEFAULT_EXPORT_SECTIONS)


def _configure_document(doc: Document, export_format: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5 if export_format == "printable" else 10)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.12

    _configure_builtin_heading_styles(doc)
    _ensure_paragraph_style(doc, "TC Low Key", RGBColor(100, 100, 100), Pt(9))
    _ensure_paragraph_style(doc, "TC Front Meta", RGBColor(72, 72, 72), Pt(10))
    _ensure_paragraph_style(doc, "TC Source", RGBColor(110, 110, 110), Pt(8.5))
    review = _ensure_paragraph_style(doc, "TC Review Flag", REVIEW_GRAY, Pt(9))
    review.paragraph_format.left_indent = Inches(0.12)
    review.paragraph_format.space_after = Pt(2)
    quote = _ensure_paragraph_style(doc, "TC Quote", RGBColor(72, 72, 72), Pt(10))
    quote.paragraph_format.left_indent = Inches(0.18)
    quote.paragraph_format.right_indent = Inches(0.08)
    quote.paragraph_format.space_before = Pt(3)
    quote.paragraph_format.space_after = Pt(5)


def _configure_builtin_heading_styles(doc: Document) -> None:
    title = doc.styles["Title"]
    _set_font(title, size=Pt(18), bold=True)
    title.paragraph_format.space_after = Pt(8)

    for name, size, before, after in [
        ("Heading 1", Pt(15), Pt(14), Pt(6)),
        ("Heading 2", Pt(12.5), Pt(9), Pt(4)),
        ("Heading 3", Pt(11), Pt(6), Pt(3)),
        ("Heading 4", Pt(10.5), Pt(5), Pt(2)),
    ]:
        style = doc.styles[name]
        _set_font(style, size=size, bold=True)
        style.paragraph_format.space_before = before
        style.paragraph_format.space_after = after


def _add_front_matter(
    doc: Document,
    course_state: dict[str, Any],
    export_format: str,
) -> None:
    source = course_state.get("source", {})
    meta = source.get("detected_meta", {})
    title = meta.get("course_title") or source.get("file") or "TextCore 导出"
    doc.add_heading(str(title), level=0)

    rows = [
        ("课程名", title),
        ("讲师", meta.get("teacher")),
        ("来源文件", source.get("file")),
        ("导出用途", "完整留档" if export_format == "archive" else "简洁可打印"),
        ("生成说明", "TextCore 根据课程转写稿整理，适合打印或归档前浏览确认。"),
    ]
    if export_format == "archive":
        rows.extend(
            [
                ("课程 ID", course_state.get("course_id")),
                ("导入时间", source.get("imported_at")),
            ]
        )

    for label, value in rows:
        if value:
            paragraph = doc.add_paragraph(style="TC Front Meta")
            paragraph.add_run(f"{label}：").bold = True
            paragraph.add_run(_plain_text(str(value)))


def _add_summary(doc: Document, course_state: dict[str, Any]) -> None:
    _add_section_heading(doc, "课程摘要")
    summary = course_state.get("global", {}).get("course_summary") or "暂无课程摘要。"
    _add_paragraphs(doc, str(summary))


def _add_version(doc: Document, version_key: str, version: dict[str, Any]) -> None:
    _add_section_heading(doc, VERSION_LABELS[version_key])
    compression = version.get("compression")
    if compression is not None:
        _add_low_key_paragraph(doc, f"压缩率：{compression:.0%}")
    _append_markdown(doc, str(version.get("body_md") or "暂无正文。"), base_level=2)


def _add_cards(doc: Document, cards: list[dict[str, Any]]) -> None:
    _add_section_heading(doc, "知识卡片")
    if not cards:
        _add_low_key_paragraph(doc, "暂无知识卡片。")
        return
    for card in cards:
        doc.add_heading(str(card.get("title") or card.get("card_id") or "知识卡片"), level=2)
        _add_paragraphs(doc, card.get("summary"))
        _add_bullets(doc, card.get("core_points", []))
        if card.get("example"):
            _add_quote(doc, f"课堂例子：{card['example']}")


def _add_materials(doc: Document, materials: list[dict[str, Any]]) -> None:
    _add_section_heading(doc, "作文素材")
    if not materials:
        _add_low_key_paragraph(doc, "暂无作文素材。")
        return
    for material in materials:
        heading = material.get("title") or material.get("material_id") or "作文素材"
        doc.add_heading(str(heading), level=2)
        if material.get("theme"):
            _add_low_key_paragraph(doc, "主题：" + "、".join(material["theme"]))
        _add_paragraphs(doc, material.get("usable_expression"))
        _add_paragraphs(doc, material.get("teacher_comment"))
        if material.get("usage_suggestion"):
            _add_quote(doc, f"用法：{material['usage_suggestion']}")


def _add_review_flags(
    doc: Document,
    flags: list[dict[str, Any]],
    export_format: str,
) -> None:
    _add_section_heading(doc, "复核标记")
    if not flags:
        _add_low_key_paragraph(doc, "暂无待复核项。")
        return
    for flag in flags:
        severity = flag.get("severity") or "medium"
        flag_id = flag.get("flag_id", "")
        text = str(flag.get("text") or "")
        suggestion = flag.get("suggestion")
        reason = flag.get("reason")

        paragraph = doc.add_paragraph(style="TC Review Flag")
        title_run = paragraph.add_run(f"{flag_id or '复核项'} [{severity}] ")
        title_run.bold = True
        paragraph.add_run(_plain_text(text))
        if suggestion:
            paragraph.add_run(f" -> {_plain_text(str(suggestion))}")
        if reason and export_format == "printable":
            paragraph.add_run(f"（{_plain_text(str(reason))}）")

        if export_format == "archive":
            detail_items = [
                f"类别：{flag.get('category')}" if flag.get("category") else "",
                f"状态：{flag.get('status')}" if flag.get("status") else "",
                f"段落：{flag.get('pid')}" if flag.get("pid") else "",
                f"分块：{flag.get('chunk_id')}" if flag.get("chunk_id") else "",
                f"原因：{reason}" if reason else "",
            ]
            detail = "；".join(_plain_text(item) for item in detail_items if item)
            if detail:
                _add_review_detail(doc, detail)


def _add_classics(
    doc: Document,
    refs: list[dict[str, Any]],
    export_format: str,
) -> None:
    _add_section_heading(doc, "古文原文与资料")
    matched_refs = [ref for ref in refs if ref.get("matched")]
    if not matched_refs:
        _add_low_key_paragraph(doc, "暂无已匹配的古文资料。")
        return
    for ref in matched_refs:
        title = ref.get("title") or ref.get("ref_id") or "古文引用"
        writer = f" - {ref['writer']}" if ref.get("writer") else ""
        doc.add_heading(f"{title}{writer}", level=2)
        _add_classics_block(doc, "原文", ref.get("canonical_text"), quote=True)
        _add_classics_block(doc, "译文", ref.get("translation"))
        _add_classics_block(doc, "注释", ref.get("remark"))
        _add_classics_block(doc, "赏析", ref.get("shangxi"))
        _add_classics_source(doc, ref, export_format)
        if export_format == "archive" and ref.get("diffs"):
            doc.add_heading("复核差异", level=3)
            for diff in ref.get("diffs", []):
                if isinstance(diff, dict):
                    raw = diff.get("raw") or ""
                    canonical = diff.get("canonical") or ""
                    pid = f"（{diff.get('pid')}）" if diff.get("pid") else ""
                    _add_review_detail(doc, f"{pid}{raw} -> {canonical}")


def _add_archive_notes(doc: Document, course_state: dict[str, Any]) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("留档信息", level=1)
    quality = course_state.get("quality", {})
    _add_low_key_paragraph(
        doc,
        "；".join(
            item
            for item in [
                (
                    f"质量分：{quality.get('quality_score')}"
                    if quality.get("quality_score") is not None
                    else ""
                ),
                f"覆盖：{quality.get('coverage')}" if quality.get("coverage") else "",
                (
                    "建议人工复核：是"
                    if quality.get("recommended_human_review")
                    else "建议人工复核：否"
                ),
            ]
            if item
        ),
    )
    for risk in quality.get("main_risks", []):
        _add_quote(doc, str(risk))


def _add_classics_block(
    doc: Document,
    label: str,
    text: Any,
    *,
    quote: bool = False,
) -> None:
    if not text:
        return
    doc.add_heading(label, level=3)
    for paragraph in str(text).splitlines():
        if paragraph.strip():
            if quote:
                _add_quote(doc, paragraph.strip())
            else:
                _add_markdown_paragraph(doc, paragraph.strip())


def _add_classics_source(
    doc: Document,
    ref: dict[str, Any],
    export_format: str,
) -> None:
    source_items = [
        f"资料库：{ref.get('source')}" if ref.get("source") else "",
        f"朝代：{ref.get('dynasty')}" if ref.get("dynasty") else "",
    ]
    if export_format == "archive":
        source_items.extend(
            [
                f"链接：{ref.get('ref_url')}" if ref.get("ref_url") else "",
                f"引用 ID：{ref.get('ref_id')}" if ref.get("ref_id") else "",
                f"分块：{ref.get('chunk_id')}" if ref.get("chunk_id") else "",
                (
                    f"置信度：{ref.get('confidence'):.0%}"
                    if isinstance(ref.get("confidence"), int | float)
                    else ""
                ),
            ]
        )
    source = "；".join(item for item in source_items if item)
    if source:
        doc.add_heading("来源", level=3)
        doc.add_paragraph(_plain_text(source), style="TC Source")


def _append_markdown(doc: Document, markdown: str, *, base_level: int) -> None:
    for raw_line in _normalize_markup(markdown).splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = _markdown_heading_level(len(heading.group(1)), base_level)
            doc.add_heading(_plain_text(heading.group(2)), level=level)
            continue
        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        if bullet:
            _add_markdown_paragraph(doc, bullet.group(1), style="List Bullet")
            continue
        numbered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if numbered:
            _add_markdown_paragraph(doc, numbered.group(1), style="List Number")
            continue
        quote = re.match(r"^\s*>\s+(.+)$", line)
        if quote:
            _add_quote(doc, quote.group(1))
            continue
        _add_markdown_paragraph(doc, line)


def _markdown_heading_level(hash_count: int, base_level: int) -> int:
    if hash_count <= 2:
        return base_level
    return min(4, base_level + hash_count - 2)


def _normalize_markup(text: str) -> str:
    normalized = str(text)
    normalized = re.sub(
        r"<h([1-6])[^>]*>(.*?)</h\1>",
        lambda match: "\n" + "#" * int(match.group(1)) + " " + match.group(2) + "\n",
        normalized,
        flags=re.IGNORECASE | re.DOTALL,
    )
    normalized = re.sub(
        r"<p[^>]*>(.*?)</p>",
        lambda match: "\n" + match.group(1) + "\n",
        normalized,
        flags=re.IGNORECASE | re.DOTALL,
    )
    normalized = re.sub(r"<br\s*/?>", "\n", normalized, flags=re.IGNORECASE)
    normalized = re.sub(r"</?(strong|b)>", "**", normalized, flags=re.IGNORECASE)
    return normalized


def _add_paragraphs(doc: Document, text: Any) -> None:
    if not text:
        return
    for paragraph in str(text).splitlines():
        if paragraph.strip():
            _add_markdown_paragraph(doc, paragraph.strip())


def _add_bullets(doc: Document, items: Iterable[Any]) -> None:
    for item in items:
        if item:
            _add_markdown_paragraph(doc, str(item), style="List Bullet")


def _add_quote(doc: Document, text: str) -> None:
    _add_markdown_paragraph(doc, text, style="TC Quote")


def _add_low_key_paragraph(doc: Document, text: str) -> None:
    if text:
        doc.add_paragraph(_plain_text(text), style="TC Low Key")


def _add_review_detail(doc: Document, text: str) -> None:
    if text:
        paragraph = doc.add_paragraph(style="TC Review Flag")
        paragraph.paragraph_format.left_indent = Inches(0.24)
        paragraph.add_run(_plain_text(text))


def _add_markdown_paragraph(doc: Document, text: Any, *, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    _add_inline_runs(paragraph, str(text))
    return paragraph


def _add_inline_runs(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+\*)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(_inline_text(text[position : match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(_inline_text(token[2:-2]))
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(_inline_text(token[1:-1]))
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        else:
            run = paragraph.add_run(_inline_text(token[1:-1]))
            run.italic = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(_inline_text(text[position:]))


def _add_section_heading(doc: Document, text: str) -> None:
    if any(paragraph.text.strip() for paragraph in doc.paragraphs):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)
    doc.add_heading(text, level=1)


def _plain_text(text: str) -> str:
    return _inline_text(text).strip()


def _inline_text(text: str) -> str:
    text = html.unescape(str(text))
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


def _ensure_paragraph_style(
    doc: Document,
    name: str,
    color: RGBColor,
    size: Pt,
):
    try:
        style = doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.color.rgb = color
    style.font.size = size
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return style


def _set_font(style, *, size: Pt, bold: bool = False) -> None:
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = size
    style.font.bold = bold
    style.font.color.rgb = RGBColor(31, 31, 31)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)
