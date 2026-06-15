from __future__ import annotations

import io

from docx import Document
from docx.shared import RGBColor

from textcore.contracts.course_state import load_example
from textcore.exporters.docx_export import export_course_docx


def test_docx_export_reopens_and_contains_selected_sections() -> None:
    state = load_example()

    content = export_course_docx(
        state,
        sections=["summary", "concise", "cards", "materials", "review"],
        export_format="printable",
    )

    doc = Document(io.BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    source_file = state["source"]["file"]
    assert "课程名：" in text
    assert "讲师：张老师" in text
    assert f"来源文件：{source_file}" in text
    assert "课程摘要" in text
    assert "精简整理" in text
    assert "知识卡片" in text
    assert "作文素材" in text
    assert "复核标记" in text
    assert "保真清洗" not in text


def test_docx_export_honors_version_and_classics_selection() -> None:
    state = load_example()

    content = export_course_docx(
        state,
        sections=["faithful", "outline", "classics"],
        export_format="archive",
    )

    doc = Document(io.BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "保真清洗" in text
    assert "结构提纲" in text
    assert "古文原文与资料" in text
    assert "精简整理" not in text
    assert "课程摘要" not in text
    assert "完整留档" in text
    assert "留档信息" in text


def test_docx_export_renders_markdown_headings_lists_quotes_and_bold() -> None:
    state = load_example()
    state["versions"]["concise"]["body_md"] = "\n".join(
        [
            "## 阅读方法",
            "这段需要**重点标记**，也要保留 `术语`。",
            "### 课堂提醒",
            "- **明线**是事件推进",
            "> **老师提醒**：先找人物动作。",
        ]
    )

    content = export_course_docx(
        state,
        sections=["concise"],
        export_format="printable",
    )

    doc = Document(io.BytesIO(content))
    paragraph_by_text = {paragraph.text: paragraph for paragraph in doc.paragraphs}
    assert paragraph_by_text["阅读方法"].style.name == "Heading 2"
    assert paragraph_by_text["课堂提醒"].style.name == "Heading 3"
    assert paragraph_by_text["明线是事件推进"].style.name == "List Bullet"
    assert paragraph_by_text["老师提醒：先找人物动作。"].style.name == "TC Quote"

    body = paragraph_by_text["这段需要重点标记，也要保留 术语。"]
    assert any(run.text == "重点标记" and run.bold for run in body.runs)
    bullet = paragraph_by_text["明线是事件推进"]
    assert any(run.text == "明线" and run.bold for run in bullet.runs)
    quote = paragraph_by_text["老师提醒：先找人物动作。"]
    assert any(run.text == "老师提醒" and run.bold for run in quote.runs)


def test_docx_export_layers_classics_blocks_and_quotes_original_text() -> None:
    state = load_example()

    content = export_course_docx(
        state,
        sections=["classics"],
        export_format="archive",
    )

    doc = Document(io.BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    for heading in ["原文", "译文", "注释", "赏析", "来源", "复核差异"]:
        assert heading in text
    assert "链接：https://www.gushiwen.cn/" in text
    assert "引用 ID：ref_001" in text
    assert "置信度：91%" in text

    original_heading_index = _paragraph_index(doc, "原文")
    original = _next_text_paragraph(doc, original_heading_index)
    assert original.style.name == "TC Quote"
    assert "醉叟者，不知何地人" in original.text

    translation_heading_index = _paragraph_index(doc, "译文")
    translation = _next_text_paragraph(doc, translation_heading_index)
    assert translation.style.name != "TC Quote"
    assert "醉叟这个人" in translation.text


def test_docx_export_makes_printable_and_archive_outputs_distinct() -> None:
    state = load_example()

    printable = export_course_docx(
        state,
        sections=["review", "classics"],
        export_format="printable",
    )
    archive = export_course_docx(
        state,
        sections=["review", "classics"],
        export_format="archive",
    )

    printable_doc = Document(io.BytesIO(printable))
    archive_doc = Document(io.BytesIO(archive))
    printable_text = "\n".join(paragraph.text for paragraph in printable_doc.paragraphs)
    archive_text = "\n".join(paragraph.text for paragraph in archive_doc.paragraphs)

    assert "导出用途：简洁可打印" in printable_text
    assert "导出用途：完整留档" in archive_text
    assert "类别：classical_typo" not in printable_text
    assert "类别：classical_typo" in archive_text
    assert "链接：https://www.gushiwen.cn/" not in printable_text
    assert "链接：https://www.gushiwen.cn/" in archive_text
    assert "留档信息" not in printable_text
    assert "留档信息" in archive_text

    review_style = printable_doc.styles["TC Review Flag"]
    assert review_style.font.color.rgb == RGBColor(104, 104, 104)


def _paragraph_index(doc: Document, text: str) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text == text:
            return index
    raise AssertionError(f"missing paragraph: {text}")


def _next_text_paragraph(doc: Document, index: int):
    for paragraph in doc.paragraphs[index + 1 :]:
        if paragraph.text.strip():
            return paragraph
    raise AssertionError("missing following text paragraph")
