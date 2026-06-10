from __future__ import annotations

import html
import json
import re
import time
from dataclasses import dataclass
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "prototype" / "full-samples.js"

SAMPLES = {
    "sample-classical": {
        "title": "五上寒假第三/四讲",
        "subtitle": "《偷钱》线索讲评 + 文言文《醉叟传》",
        "source": ROOT
        / "素材"
        / "五上-人文综合涵养-寒假-第三讲-隐显-偷钱+第四讲-文言文-醉叟传1.docx",
        "max_chars": 3600,
    },
    "sample-essay": {
        "title": "五上寒假第七讲",
        "subtitle": "现代文阅读 + 作文点评",
        "source": ROOT / "素材" / "五上-人文综合涵养-寒假-第七讲-阅读理解+作文点评2.docx",
        "max_chars": 3600,
    },
}


@dataclass
class Chunk:
    chunk_id: str
    title: str
    raw_chars: int
    paragraphs: list[str]
    start_para: int
    end_para: int


def read_paragraphs(path: Path) -> list[str]:
    doc = Document(path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def is_timestamp(text: str) -> bool:
    return bool(re.fullmatch(r"[\u4e00-\u9fa5A-Za-z_·\s]{1,24}\s+\d{2}:\d{2}:\d{2}", text))


def clean_text(text: str) -> str:
    text = re.sub(r"^[\u4e00-\u9fa5A-Za-z_·\s]{1,24}\s+\d{2}:\d{2}:\d{2}\s*", "", text)
    replacements = [
        (r"\bMm+\b|hmm|yes", ""),
        (r"[嗯呃]{1,3}[，。 ]*", ""),
        (r"呵呵|哈哈", ""),
        (r"是不是(啊|呀)?[？?，,。 ]*", ""),
        (r"对吧[？?，,。 ]*", ""),
        (r"能理解吧[？?，,。 ]*", ""),
        (r"懂了吗[？?，,。 ]*", ""),
        (r"好[，,。 ]*(?=第|接下来|我们|那么|来|翻|往下)", ""),
        (r"这个这个|那个那个|就就|把把|的的的", ""),
        (r"\s+", " "),
    ]
    for pattern, repl in replacements:
        text = re.sub(pattern, repl, text, flags=re.IGNORECASE)
    text = text.replace("罪首", "醉叟").replace("最手", "醉叟").replace("醉手", "醉叟")
    text = text.replace("古兽行销", "骨瘦形销")
    text = re.sub(r"([。！？；])\1+", r"\1", text)
    text = re.sub(r"[，,]\s*[，,]+", "，", text)
    return text.strip(" ，,。")


def sentence_split(text: str) -> list[str]:
    parts = re.split(r"(?<=[。！？；])", text)
    return [p.strip() for p in parts if len(p.strip()) >= 8]


def chunk_title(sample_id: str, text: str, index: int) -> str:
    rules = []
    if sample_id == "sample-classical":
        rules = [
            ("明线" in text or "暗线" in text, "线索、明线与暗线"),
            ("偷钱" in text and ("母亲" in text or "父亲" in text), "《偷钱》情节与人物心理"),
            ("祖母" in text or "宽容" in text or "尊严" in text, "犯错后的家庭教育"),
            ("犯过" in text or "严重的错误" in text or "分享" in text, "课堂拓展：错误与反思"),
            ("袁宏道" in text or "公安派" in text or "文言文" in text, "进入《醉叟传》与作者背景"),
            ("醉叟者" in text or "姓字" in text or "呼曰" in text, "《醉叟传》第一段句读与翻译"),
            ("蜈蚣" in text or "黄竹篮" in text or "市儿" in text, "醉叟形象细节与复核点"),
        ]
    else:
        rules = [
            ("三步走" in text or "题型" in text, "现代文阅读三步法"),
            ("放学路" in text or "儿歌" in text, "《放学路上》的开头与主题"),
            ("概括题" in text or "答题区间" in text, "概括题：区间、关键词与格式"),
            ("引用" in text or "赏析题" in text or "表达效果" in text, "引用作用与赏析题"),
            ("豪华笼子" in text or "贵重行李" in text or "比喻" in text, "比喻信息与中心理解"),
            ("W教授" in text or "克隆" in text, "小说讲评：W教授与科学反思"),
            ("作文" in text or "初春" in text or "残雪" in text or "显得" in text, "作文点评：写景语言更新"),
        ]
    for matched, title in rules:
        if matched:
            return title
    return f"课堂片段 {index}"


def segment(sample_id: str, paragraphs: list[str], max_chars: int) -> list[Chunk]:
    chunks: list[Chunk] = []
    current: list[str] = []
    start_para = 1
    current_chars = 0
    for idx, raw in enumerate(paragraphs, start=1):
        if idx == 1:
            continue
        if is_timestamp(raw):
            continue
        cleaned = clean_text(raw)
        if not cleaned:
            continue
        is_noise = len(cleaned) > 220 and len(re.findall(r"[。！？]", cleaned)) <= 1 and cleaned.count("我") > 12
        if is_noise:
            cleaned = f"[转写噪声片段，建议复核或删除] {cleaned[:180]}..."
        if current and current_chars + len(cleaned) > max_chars:
            body = "\n".join(current)
            chunks.append(
                Chunk(
                    chunk_id=f"c{len(chunks) + 1:02d}",
                    title=chunk_title(sample_id, body, len(chunks) + 1),
                    raw_chars=current_chars,
                    paragraphs=current,
                    start_para=start_para,
                    end_para=idx - 1,
                )
            )
            current = []
            start_para = idx
            current_chars = 0
        current.append(cleaned)
        current_chars += len(cleaned)
    if current:
        body = "\n".join(current)
        chunks.append(
            Chunk(
                chunk_id=f"c{len(chunks) + 1:02d}",
                title=chunk_title(sample_id, body, len(chunks) + 1),
                raw_chars=current_chars,
                paragraphs=current,
                start_para=start_para,
                end_para=len(paragraphs),
            )
        )
    return chunks


def important_sentences(paragraphs: list[str], limit: int = 8) -> list[str]:
    keywords = [
        "第一",
        "第二",
        "第三",
        "重点",
        "所以",
        "说明",
        "意味着",
        "作用",
        "方法",
        "题型",
        "明线",
        "暗线",
        "母亲",
        "父亲",
        "祖母",
        "宽容",
        "醉叟",
        "袁宏道",
        "概括题",
        "赏析题",
        "引用",
        "比喻",
        "作文",
    ]
    scored: list[tuple[int, str]] = []
    for para in paragraphs:
        for sent in sentence_split(para):
            score = sum(2 for kw in keywords if kw in sent)
            if 18 <= len(sent) <= 170:
                score += 1
            if sent.startswith("[转写噪声"):
                score -= 3
            if score > 0:
                scored.append((score, sent))
    selected: list[str] = []
    seen = set()
    for _, sent in sorted(scored, key=lambda item: item[0], reverse=True):
        key = sent[:24]
        if key in seen:
            continue
        seen.add(key)
        selected.append(sent)
        if len(selected) >= limit:
            break
    if len(selected) < 3:
        for para in paragraphs:
            if para.startswith("[转写噪声"):
                continue
            selected.append(para[:150])
            if len(selected) >= limit:
                break
    return selected


def review_flags(sample_id: str, chunks: list[Chunk]) -> list[str]:
    joined = "\n".join("\n".join(c.paragraphs) for c in chunks)
    flags = []
    if sample_id == "sample-classical":
        flags.extend(
            [
                "“醉叟”在原转写中多次出现为“罪首/最手/醉手”，本次按语境统一校为“醉叟”，需核对讲义。",
                "“尽日酣沉”疑似被转写为“近日酣沉”，需核对原文。",
                "“京、离”等地名和《醉叟传》原句需按讲义或古文原文复核。",
            ]
        )
    else:
        flags.extend(
            [
                "《放学路上》篇名疑似转写为“萧氏的放学路上”，需核对讲义。",
                "“骨瘦如柴 / 骨瘦形销”相关作文点评片段存在转写混乱，需复核。",
                "尾段出现大量课堂杂音和无意义转写，学习整理版默认删除。",
            ]
        )
    if "转写噪声片段" in joined:
        flags.append("全文中存在疑似 ASR 噪声块，已在保真版保留低干扰标记。")
    return flags


def chunk_footer(index: int, chunks: list[Chunk]) -> str:
    prev_chunk = chunks[index - 1].chunk_id if index > 0 else ""
    next_chunk = chunks[index + 1].chunk_id if index < len(chunks) - 1 else ""
    return f"""
      <nav class="chunk-footer" aria-label="片段跳转">
        <button class="tiny-button" data-jump-chunk="{prev_chunk}" {"disabled" if not prev_chunk else ""}>上一段</button>
        <button class="tiny-button" data-jump-toc>回到目录</button>
        <button class="tiny-button" data-jump-chunk="{next_chunk}" {"disabled" if not next_chunk else ""}>下一段</button>
      </nav>
    """


def render_clean(chunks: list[Chunk]) -> str:
    sections = []
    for index, chunk in enumerate(chunks):
        paras = "\n".join(f"<p>{html.escape(p)}</p>" for p in chunk.paragraphs)
        sections.append(
            f"""
            <section class="long-section" id="{chunk.chunk_id}">
              <h2>{html.escape(chunk.title)}</h2>
              <p class="chunk-meta">片段 {chunk.chunk_id}｜原段落 {chunk.start_para}-{chunk.end_para}｜约 {chunk.raw_chars} 字</p>
              {paras}
              {chunk_footer(index, chunks)}
            </section>
            """
        )
    return "\n".join(sections)


def render_study(chunks: list[Chunk]) -> str:
    sections = []
    for index, chunk in enumerate(chunks):
        points = important_sentences(chunk.paragraphs, 7)
        li = "\n".join(f"<li>{html.escape(p)}</li>" for p in points)
        sections.append(
            f"""
            <section class="long-section" id="{chunk.chunk_id}">
              <h2>{html.escape(chunk.title)}</h2>
              <p class="chunk-meta">学习整理｜覆盖原段落 {chunk.start_para}-{chunk.end_para}｜保留重点 {len(points)} 条</p>
              <ul>{li}</ul>
              {chunk_footer(index, chunks)}
            </section>
            """
        )
    return "\n".join(sections)


def render_outline(chunks: list[Chunk]) -> str:
    items = []
    for chunk in chunks:
        points = important_sentences(chunk.paragraphs, 4)
        sub = "".join(f"<li>{html.escape(p[:88])}</li>" for p in points)
        items.append(
            f"""
            <li>
              <strong>{html.escape(chunk.title)}</strong>
              <span class="chunk-meta">片段 {chunk.chunk_id}｜原段落 {chunk.start_para}-{chunk.end_para}</span>
              <ul>{sub}</ul>
            </li>
            """
        )
    return f'<section class="long-section"><h2>全文结构提纲</h2><ol class="outline-list">{"".join(items)}</ol></section>'


def build_sample(sample_id: str, config: dict) -> dict:
    started = time.perf_counter()
    raw_paragraphs = read_paragraphs(config["source"])
    extraction_done = time.perf_counter()
    chunks = segment(sample_id, raw_paragraphs, config["max_chars"])
    segmentation_done = time.perf_counter()
    versions = {
        "clean": render_clean(chunks),
        "study": render_study(chunks),
        "outline": render_outline(chunks),
    }
    generation_done = time.perf_counter()
    clean_chars = sum(c.raw_chars for c in chunks)
    return {
        "title": config["title"],
        "subtitle": config["subtitle"],
        "sourceFile": config["source"].name,
        "stats": {
            "paragraphs": len(raw_paragraphs),
            "rawChars": sum(len(p) for p in raw_paragraphs),
            "cleanChars": clean_chars,
            "chunks": len(chunks),
            "chunkTargetChars": config["max_chars"],
        },
        "timing": {
            "mode": "local_rule_based_prototype",
            "totalSeconds": round(generation_done - started, 2),
            "extractSeconds": round(extraction_done - started, 2),
            "segmentSeconds": round(segmentation_done - extraction_done, 2),
            "renderSeconds": round(generation_done - segmentation_done, 2),
            "llmEstimate": "正式高质量 LLM 分块处理预计 6-15 分钟/篇，取决于模型、并发和质量检查轮次。",
        },
        "chunks": [
            {
                "id": c.chunk_id,
                "title": c.title,
                "startPara": c.start_para,
                "endPara": c.end_para,
                "chars": c.raw_chars,
            }
            for c in chunks
        ],
        "reviewItems": review_flags(sample_id, chunks),
        "versions": versions,
    }


def main() -> None:
    built = {sample_id: build_sample(sample_id, config) for sample_id, config in SAMPLES.items()}
    payload = json.dumps(built, ensure_ascii=False, indent=2)
    OUT.write_text(f"window.fullLessonSamples = {payload};\n", encoding="utf-8")
    print(f"wrote {OUT}")
    for sample_id, sample in built.items():
        print(
            sample_id,
            f"paras={sample['stats']['paragraphs']}",
            f"raw={sample['stats']['rawChars']}",
            f"chunks={sample['stats']['chunks']}",
            f"seconds={sample['timing']['totalSeconds']}",
        )


if __name__ == "__main__":
    main()
